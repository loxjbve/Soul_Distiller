from __future__ import annotations

import io

import pytest
from docx import Document
from pypdf import PdfWriter

from app.pipeline.extractors import UnsupportedScannedPdfError, extract_text


def build_text_pdf_bytes(text: str) -> bytes:
    stream = f"BT\n/F1 24 Tf\n72 72 Td\n({text}) Tj\nET"
    objects = [
        "<< /Type /Catalog /Pages 2 0 R >>",
        "<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        "<< /Type /Page /Parent 2 0 R /MediaBox [0 0 300 144] /Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>",
        f"<< /Length {len(stream.encode('latin-1'))} >>\nstream\n{stream}\nendstream",
        "<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]
    chunks = [b"%PDF-1.4\n"]
    offsets = [0]
    for index, obj in enumerate(objects, start=1):
        offsets.append(sum(len(chunk) for chunk in chunks))
        chunks.append(f"{index} 0 obj\n{obj}\nendobj\n".encode("latin-1"))
    xref_start = sum(len(chunk) for chunk in chunks)
    chunks.append(f"xref\n0 {len(objects) + 1}\n".encode("latin-1"))
    chunks.append(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        chunks.append(f"{offset:010d} 00000 n \n".encode("latin-1"))
    chunks.append(f"trailer\n<< /Root 1 0 R /Size {len(objects) + 1} >>\n".encode("latin-1"))
    chunks.append(f"startxref\n{xref_start}\n%%EOF\n".encode("latin-1"))
    return b"".join(chunks)


def test_extract_html_removes_script():
    html = b"""
    <html><head><title>Profile</title><script>alert('x')</script></head>
    <body><h1>Alice</h1><p>Hello world.</p></body></html>
    """
    result = extract_text("profile.html", html)
    assert result.title == "Profile"
    assert "alert" not in result.clean_text
    assert "Alice" in result.clean_text


def test_extract_json_flattens_nested_content():
    payload = b'{"profile":{"name":"Alice","bio":"Writes essays"},"id":"skip-me"}'
    result = extract_text("profile.json", payload)
    assert "profile.name: Alice" in result.clean_text
    assert "skip-me" not in result.clean_text


def test_extract_jsonl_flattens_each_line():
    payload = b'{"name":"Alice"}\n{"name":"Bob"}\nNot a JSON line\n{"bio":"Writer"}'
    result = extract_text("data.jsonl", payload)
    assert "name: Alice" in result.clean_text
    assert "name: Bob" in result.clean_text
    assert "Not a JSON line" in result.clean_text
    assert "bio: Writer" in result.clean_text
    assert result.metadata["format"] == "jsonl"
    assert len(result.segments) == 4


def test_extract_docx_reads_paragraphs_and_tables():
    doc = Document()
    doc.add_heading("Alice Memo", level=1)
    doc.add_paragraph("First paragraph.")
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Role"
    table.cell(0, 1).text = "Writer"
    buffer = io.BytesIO()
    doc.save(buffer)
    result = extract_text("memo.docx", buffer.getvalue())
    assert result.title == "Alice Memo"
    assert "First paragraph." in result.clean_text
    assert "Role | Writer" in result.clean_text


def test_extract_pdf_text_keeps_page_metadata():
    pdf_bytes = build_text_pdf_bytes("Hello PDF")
    result = extract_text("sample.pdf", pdf_bytes)
    assert result.metadata["page_count"] == 1
    assert result.segments[0].metadata["page_number"] == 1
    assert "Hello PDF" in result.clean_text


def test_extract_blank_pdf_raises_scanned_error():
    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    buffer = io.BytesIO()
    writer.write(buffer)
    with pytest.raises(UnsupportedScannedPdfError):
        extract_text("blank.pdf", buffer.getvalue())
