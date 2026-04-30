from __future__ import annotations

import csv
import io
import json
import re
from pathlib import Path
from typing import Any

from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from pypdf import PdfReader

from app.persona_plus import is_persona_mode
from app.schemas import ExtractedSegment, ExtractionResult
from app.utils.text import guess_language, normalize_whitespace


class ExtractionError(Exception):
    """Base extraction error."""


class UnsupportedDocumentError(ExtractionError):
    """Raised when a file extension is not supported."""


class UnsupportedScannedPdfError(ExtractionError):
    """Raised when PDF appears to be image-only."""


MEANINGLESS_JSON_KEYS = {"id", "uuid", "hash", "etag", "checksum", "token"}
PERSONA_INTAKE_KEYS = {"channel", "date", "approx_date", "confidence", "scope"}
PERSONA_SENDER_KEYS = ("sender", "from", "author", "user", "username", "name", "speaker")
PERSONA_TEXT_KEYS = ("text", "content", "body", "message", "caption")
PERSONA_TIME_KEYS = ("timestamp", "date", "time", "created_at", "sent_at", "datetime")
PERSONA_ID_KEYS = ("id", "message_id", "msg_id", "uid")
PERSONA_COMMENT_PATTERN = re.compile(r"<!--\s*(?P<body>.*?)\s*-->", re.DOTALL)
PERSONA_BRACKET_CHAT_PATTERN = re.compile(r"^\[(?P<date>[^\]]+)\]\s*(?P<speaker>[^:：]{1,80})[:：]\s*(?P<text>.+)$")
PERSONA_DATED_CHAT_PATTERN = re.compile(
    r"^(?P<date>\d{4}[-/年]\d{1,2}(?:[-/月]\d{1,2})?(?:[ T]\d{1,2}:\d{2}(?::\d{2})?)?)\s+"
    r"(?P<speaker>[^:：]{1,80})[:：]\s*(?P<text>.+)$"
)
PERSONA_SIMPLE_CHAT_PATTERN = re.compile(r"^(?P<speaker>[^:：]{1,80})[:：]\s*(?P<text>.+)$")


def extract_text(filename: str, content: bytes, *, mode: str | None = None) -> ExtractionResult:
    extension = Path(filename).suffix.lower()
    persona_mode = is_persona_mode(mode)
    if extension in {".html", ".htm"}:
        return _extract_html(content)
    if extension == ".json":
        if persona_mode:
            return _extract_persona_json(content)
        return _extract_json(content)
    if extension == ".jsonl":
        if persona_mode:
            return _extract_persona_jsonl(content)
        return _extract_jsonl(content)
    if extension == ".csv":
        if persona_mode:
            return _extract_persona_csv(content)
        return _extract_csv(content)
    if extension == ".docx":
        return _extract_docx(content)
    if extension == ".pdf":
        return _extract_pdf(content)
    if extension in {".txt", ".md", ".log"}:
        if persona_mode:
            return _extract_persona_textual(content, source_format=extension.lstrip(".") or "text")
        return _extract_textual(content)
    raise UnsupportedDocumentError(f"Unsupported document type: {extension}")


def _extract_html(content: bytes) -> ExtractionResult:
    decoded = _decode_text(content)
    soup = BeautifulSoup(decoded, "html.parser")
    for node in soup(["script", "style", "noscript"]):
        node.decompose()
    blocks: list[str] = []
    for tag in soup.find_all(["h1", "h2", "h3", "p", "li", "blockquote"]):
        text = tag.get_text(" ", strip=True)
        if text:
            blocks.append(text)
    title = soup.title.get_text(strip=True) if soup.title else None
    raw_text = "\n".join(blocks) or soup.get_text("\n", strip=True)
    segments = [ExtractedSegment(text=normalize_whitespace(block), metadata={}) for block in blocks if block.strip()]
    if not segments and raw_text.strip():
        segments = [ExtractedSegment(text=normalize_whitespace(raw_text), metadata={})]
    clean_text = normalize_whitespace("\n\n".join(segment.text for segment in segments))
    return ExtractionResult(
        raw_text=raw_text,
        clean_text=clean_text,
        title=title,
        author_guess=None,
        created_at_guess=None,
        language=guess_language(clean_text),
        metadata={"block_count": len(segments), "format": "html"},
        segments=segments,
    )


def _extract_json(content: bytes) -> ExtractionResult:
    decoded = _decode_text(content)
    payload = json.loads(decoded)
    lines: list[str] = []
    _flatten_json(payload, [], lines)
    raw_text = "\n".join(lines)
    clean_text = normalize_whitespace(raw_text)
    segments = [ExtractedSegment(text=line, metadata={}) for line in lines if line.strip()]
    return ExtractionResult(
        raw_text=raw_text,
        clean_text=clean_text,
        title=None,
        author_guess=None,
        created_at_guess=None,
        language=guess_language(clean_text),
        metadata={"line_count": len(lines), "format": "json"},
        segments=segments,
    )


def _extract_jsonl(content: bytes) -> ExtractionResult:
    decoded = _decode_text(content)
    all_extracted_lines: list[str] = []
    for line in decoded.splitlines():
        line = line.strip()
        if not line:
            continue
        try:
            payload = json.loads(line)
            record_lines: list[str] = []
            _flatten_json(payload, [], record_lines)
            all_extracted_lines.extend(record_lines)
        except json.JSONDecodeError:
            # Fallback to plain text if a line is not valid JSON
            all_extracted_lines.append(line)

    raw_text = "\n".join(all_extracted_lines)
    clean_text = normalize_whitespace(raw_text)
    segments = [ExtractedSegment(text=line, metadata={}) for line in all_extracted_lines if line.strip()]
    return ExtractionResult(
        raw_text=raw_text,
        clean_text=clean_text,
        title=None,
        author_guess=None,
        created_at_guess=None,
        language=guess_language(clean_text),
        metadata={"line_count": len(all_extracted_lines), "format": "jsonl"},
        segments=segments,
    )


def _extract_csv(content: bytes) -> ExtractionResult:
    decoded = _decode_text(content)
    rows = _read_csv_rows(decoded)
    lines: list[str] = []
    for row in rows:
        if isinstance(row, dict):
            line = " | ".join(f"{key}: {value}" for key, value in row.items() if str(value or "").strip())
        else:
            line = " | ".join(str(value).strip() for value in row if str(value or "").strip())
        if line.strip():
            lines.append(normalize_whitespace(line))
    raw_text = "\n".join(lines)
    clean_text = normalize_whitespace(raw_text)
    segments = [ExtractedSegment(text=line, metadata={}) for line in lines if line.strip()]
    return ExtractionResult(
        raw_text=raw_text,
        clean_text=clean_text,
        title=None,
        author_guess=None,
        created_at_guess=None,
        language=guess_language(clean_text),
        metadata={"row_count": len(rows), "line_count": len(lines), "format": "csv"},
        segments=segments,
    )


def _extract_docx(content: bytes) -> ExtractionResult:
    doc = DocxDocument(io.BytesIO(content))
    parts: list[str] = []
    title = doc.core_properties.title or None
    author = doc.core_properties.author or None
    created_at_guess = None
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)
            if not title and paragraph.style and "heading" in paragraph.style.name.lower():
                title = text
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                rows.append(" | ".join(cells))
        if rows:
            parts.extend(rows)
    raw_text = "\n".join(parts)
    segments = [ExtractedSegment(text=normalize_whitespace(part), metadata={}) for part in parts if part.strip()]
    clean_text = normalize_whitespace("\n\n".join(segment.text for segment in segments))
    return ExtractionResult(
        raw_text=raw_text,
        clean_text=clean_text,
        title=title,
        author_guess=author,
        created_at_guess=created_at_guess,
        language=guess_language(clean_text),
        metadata={"part_count": len(segments), "format": "docx"},
        segments=segments,
    )


def _extract_pdf(content: bytes) -> ExtractionResult:
    reader = PdfReader(io.BytesIO(content))
    page_segments: list[ExtractedSegment] = []
    for index, page in enumerate(reader.pages, start=1):
        text = normalize_whitespace(page.extract_text() or "")
        if text:
            page_segments.append(ExtractedSegment(text=text, metadata={"page_number": index}))
    if not page_segments:
        raise UnsupportedScannedPdfError("PDF contains no extractable text. OCR is not supported in v1.")
    raw_text = "\n\n".join(segment.text for segment in page_segments)
    clean_text = normalize_whitespace(raw_text)
    metadata = {"page_count": len(reader.pages), "format": "pdf", "digital_text_pages": len(page_segments)}
    return ExtractionResult(
        raw_text=raw_text,
        clean_text=clean_text,
        title=None,
        author_guess=None,
        created_at_guess=None,
        language=guess_language(clean_text),
        metadata=metadata,
        segments=page_segments,
    )


def _extract_textual(content: bytes) -> ExtractionResult:
    decoded = _decode_text(content)
    raw_text = decoded
    paragraphs = [normalize_whitespace(chunk) for chunk in re.split(r"\n\s*\n", decoded) if chunk.strip()]
    if not paragraphs:
        paragraphs = [normalize_whitespace(decoded)]
    clean_text = normalize_whitespace("\n\n".join(paragraphs))
    segments = [ExtractedSegment(text=paragraph, metadata={}) for paragraph in paragraphs if paragraph.strip()]
    return ExtractionResult(
        raw_text=raw_text,
        clean_text=clean_text,
        title=paragraphs[0][:80] if paragraphs else None,
        author_guess=None,
        created_at_guess=None,
        language=guess_language(clean_text),
        metadata={"paragraph_count": len(segments), "format": "text"},
        segments=segments,
    )


def _extract_persona_textual(content: bytes, *, source_format: str) -> ExtractionResult:
    decoded = _decode_text(content)
    segments = _persona_segments_from_text(decoded, base_metadata={"source_format": source_format})
    if not segments:
        fallback = normalize_whitespace(decoded)
        segments = [ExtractedSegment(text=fallback, metadata={"source_format": source_format})] if fallback else []
    raw_text = decoded
    clean_text = normalize_whitespace("\n\n".join(segment.text for segment in segments))
    return ExtractionResult(
        raw_text=raw_text,
        clean_text=clean_text,
        title=segments[0].text[:80] if segments else None,
        author_guess=None,
        created_at_guess=_first_segment_date(segments),
        language=guess_language(clean_text),
        metadata={
            "paragraph_count": len(segments),
            "message_count": sum(1 for segment in segments if segment.metadata.get("speaker")),
            "format": source_format,
            "persona_plus": True,
        },
        segments=segments,
    )


def _extract_persona_json(content: bytes) -> ExtractionResult:
    decoded = _decode_text(content)
    payload = json.loads(decoded)
    segments = _persona_segments_from_json_payload(payload, source_format="json")
    if not segments:
        result = _extract_json(content)
        return ExtractionResult(
            raw_text=result.raw_text,
            clean_text=result.clean_text,
            title=result.title,
            author_guess=result.author_guess,
            created_at_guess=result.created_at_guess,
            language=result.language,
            metadata={**dict(result.metadata or {}), "persona_plus": True},
            segments=[
                ExtractedSegment(text=segment.text, metadata={**dict(segment.metadata or {}), "source_format": "json"})
                for segment in result.segments
            ],
        )
    raw_text = "\n".join(segment.text for segment in segments)
    clean_text = normalize_whitespace(raw_text)
    return ExtractionResult(
        raw_text=raw_text,
        clean_text=clean_text,
        title=_persona_payload_title(payload),
        author_guess=None,
        created_at_guess=_first_segment_date(segments),
        language=guess_language(clean_text),
        metadata={"message_count": len(segments), "format": "json", "persona_plus": True},
        segments=segments,
    )


def _extract_persona_jsonl(content: bytes) -> ExtractionResult:
    decoded = _decode_text(content)
    segments: list[ExtractedSegment] = []
    for index, line in enumerate(decoded.splitlines(), start=1):
        raw_line = line.strip()
        if not raw_line:
            continue
        try:
            payload = json.loads(raw_line)
        except json.JSONDecodeError:
            text = normalize_whitespace(raw_line)
            if text:
                segments.append(ExtractedSegment(text=text, metadata={"source_format": "jsonl", "record_index": index}))
            continue
        if isinstance(payload, dict):
            segment = _persona_segment_from_record(payload, source_format="jsonl", record_index=index)
            if segment:
                segments.append(segment)
                continue
            record_lines: list[str] = []
            _flatten_json(payload, [], record_lines)
            text = normalize_whitespace("\n".join(record_lines))
            if text:
                segments.append(ExtractedSegment(text=text, metadata={"source_format": "jsonl", "record_index": index}))
        else:
            text = normalize_whitespace(_persona_value_to_text(payload))
            if text:
                segments.append(ExtractedSegment(text=text, metadata={"source_format": "jsonl", "record_index": index}))
    raw_text = "\n".join(segment.text for segment in segments)
    clean_text = normalize_whitespace(raw_text)
    return ExtractionResult(
        raw_text=raw_text,
        clean_text=clean_text,
        title=None,
        author_guess=None,
        created_at_guess=_first_segment_date(segments),
        language=guess_language(clean_text),
        metadata={"line_count": len(segments), "format": "jsonl", "persona_plus": True},
        segments=segments,
    )


def _extract_persona_csv(content: bytes) -> ExtractionResult:
    decoded = _decode_text(content)
    rows = _read_csv_rows(decoded)
    segments: list[ExtractedSegment] = []
    for index, row in enumerate(rows, start=1):
        if not isinstance(row, dict):
            text = normalize_whitespace(" | ".join(str(value).strip() for value in row if str(value or "").strip()))
            if text:
                segments.append(ExtractedSegment(text=text, metadata={"source_format": "csv", "row_index": index}))
            continue
        segment = _persona_segment_from_record(row, source_format="csv", record_index=index)
        if segment:
            metadata = dict(segment.metadata or {})
            metadata["row_index"] = index
            segments.append(ExtractedSegment(text=segment.text, metadata=metadata))
            continue
        text = normalize_whitespace(" | ".join(f"{key}: {value}" for key, value in row.items() if str(value or "").strip()))
        if text:
            segments.append(ExtractedSegment(text=text, metadata={"source_format": "csv", "row_index": index}))
    raw_text = "\n".join(segment.text for segment in segments)
    clean_text = normalize_whitespace(raw_text)
    return ExtractionResult(
        raw_text=raw_text,
        clean_text=clean_text,
        title=None,
        author_guess=None,
        created_at_guess=_first_segment_date(segments),
        language=guess_language(clean_text),
        metadata={"row_count": len(rows), "message_count": len(segments), "format": "csv", "persona_plus": True},
        segments=segments,
    )


def _decode_text(content: bytes) -> str:
    encodings = ["utf-8", "utf-8-sig", "gb18030", "big5", "latin-1"]
    for encoding in encodings:
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="ignore")


def _read_csv_rows(decoded: str) -> list[Any]:
    stream = io.StringIO(decoded)
    sample = decoded[:4096]
    try:
        dialect = csv.Sniffer().sniff(sample) if sample.strip() else csv.excel
    except csv.Error:
        dialect = csv.excel
    try:
        has_header = csv.Sniffer().has_header(sample) if sample.strip() else False
    except csv.Error:
        has_header = True
    stream.seek(0)
    if has_header:
        reader = csv.DictReader(stream, dialect=dialect)
        return [
            {str(key or "").strip(): str(value or "").strip() for key, value in row.items() if str(key or "").strip()}
            for row in reader
        ]
    reader = csv.reader(stream, dialect=dialect)
    return [[str(value or "").strip() for value in row] for row in reader]


def _persona_segments_from_json_payload(payload: Any, *, source_format: str) -> list[ExtractedSegment]:
    records = _persona_message_records(payload)
    segments: list[ExtractedSegment] = []
    for index, record in enumerate(records, start=1):
        segment = _persona_segment_from_record(record, source_format=source_format, record_index=index)
        if segment:
            segments.append(segment)
    return segments


def _persona_message_records(payload: Any) -> list[dict[str, Any]]:
    if isinstance(payload, dict):
        messages = payload.get("messages")
        if isinstance(messages, list):
            return [item for item in messages if isinstance(item, dict)]
        for key in ("items", "records", "data", "chats"):
            value = payload.get(key)
            if isinstance(value, list) and any(isinstance(item, dict) and _persona_record_text(item) for item in value):
                return [item for item in value if isinstance(item, dict)]
        if _persona_record_text(payload):
            return [payload]
        collected: list[dict[str, Any]] = []
        _collect_persona_message_records(payload, collected)
        return collected
    if isinstance(payload, list):
        if all(isinstance(item, dict) for item in payload):
            return [item for item in payload if isinstance(item, dict)]
        collected: list[dict[str, Any]] = []
        _collect_persona_message_records(payload, collected)
        return collected
    return []


def _collect_persona_message_records(value: Any, records: list[dict[str, Any]]) -> None:
    if isinstance(value, dict):
        if _persona_record_text(value):
            records.append(value)
            return
        for child in value.values():
            _collect_persona_message_records(child, records)
        return
    if isinstance(value, list):
        for item in value[:500]:
            _collect_persona_message_records(item, records)


def _persona_segment_from_record(
    record: dict[str, Any],
    *,
    source_format: str,
    record_index: int,
) -> ExtractedSegment | None:
    text = _persona_record_text(record)
    if not text:
        return None
    metadata = _persona_metadata_from_record(record)
    metadata.setdefault("source_format", source_format)
    metadata.setdefault("record_index", record_index)
    speaker = str(metadata.get("speaker") or "").strip()
    content = f"{speaker}: {text}" if speaker else text
    return ExtractedSegment(text=normalize_whitespace(content), metadata=metadata)


def _persona_record_text(record: dict[str, Any]) -> str:
    lower = {str(key).strip().lower(): value for key, value in record.items()}
    for key in PERSONA_TEXT_KEYS:
        if key in lower:
            text = normalize_whitespace(_persona_value_to_text(lower[key]))
            if text:
                return text
    return ""


def _persona_metadata_from_record(record: dict[str, Any]) -> dict[str, Any]:
    lower = {str(key).strip().lower(): value for key, value in record.items()}
    metadata: dict[str, Any] = {}
    for key in PERSONA_INTAKE_KEYS:
        if key in lower and str(lower[key] or "").strip():
            metadata[key] = str(lower[key]).strip()
    speaker = _first_present(lower, PERSONA_SENDER_KEYS)
    if speaker:
        metadata["speaker"] = speaker
        metadata["sender_name"] = speaker
    sent_at = _first_present(lower, PERSONA_TIME_KEYS)
    if sent_at:
        metadata["date"] = metadata.get("date") or sent_at
        metadata["sent_at"] = sent_at
    message_id = _first_present(lower, PERSONA_ID_KEYS)
    if message_id:
        metadata["message_id"] = message_id
    metadata["persona_plus"] = True
    return _clean_metadata(metadata)


def _first_present(mapping: dict[str, Any], keys: tuple[str, ...]) -> str | None:
    for key in keys:
        if key in mapping:
            value = mapping[key]
            if isinstance(value, (dict, list)):
                value = _persona_value_to_text(value)
            text = str(value or "").strip()
            if text:
                return text
    return None


def _persona_segments_from_text(decoded: str, *, base_metadata: dict[str, Any] | None = None) -> list[ExtractedSegment]:
    segments: list[ExtractedSegment] = []
    current_metadata = _clean_metadata({**dict(base_metadata or {}), "persona_plus": True})
    position = 0
    for match in PERSONA_COMMENT_PATTERN.finditer(decoded):
        _append_persona_text_block_segments(decoded[position:match.start()], current_metadata, segments)
        parsed = _parse_intake_comment(match.group("body"))
        if parsed:
            current_metadata = _clean_metadata({**current_metadata, **parsed, "persona_plus": True})
        position = match.end()
    _append_persona_text_block_segments(decoded[position:], current_metadata, segments)
    return segments


def _append_persona_text_block_segments(
    text: str,
    metadata: dict[str, Any],
    segments: list[ExtractedSegment],
) -> None:
    for paragraph in [chunk for chunk in re.split(r"\n\s*\n", text) if chunk.strip()]:
        lines = [line.strip() for line in paragraph.splitlines() if line.strip()]
        parsed_lines = [_parse_persona_chat_line(line) for line in lines]
        chat_line_count = sum(1 for item in parsed_lines if item is not None)
        if len(lines) >= 2 and chat_line_count >= max(2, len(lines) // 2):
            for line_index, parsed in enumerate(parsed_lines, start=1):
                if not parsed:
                    continue
                line_metadata = _clean_metadata({**metadata, **parsed["metadata"], "line_index": line_index})
                segments.append(ExtractedSegment(text=parsed["text"], metadata=line_metadata))
            continue
        clean = normalize_whitespace(paragraph)
        if clean:
            segments.append(ExtractedSegment(text=clean, metadata=_clean_metadata(metadata)))


def _parse_persona_chat_line(line: str) -> dict[str, Any] | None:
    for pattern in (PERSONA_BRACKET_CHAT_PATTERN, PERSONA_DATED_CHAT_PATTERN, PERSONA_SIMPLE_CHAT_PATTERN):
        match = pattern.match(line)
        if not match:
            continue
        speaker = str(match.groupdict().get("speaker") or "").strip()
        text = normalize_whitespace(match.groupdict().get("text") or "")
        if not speaker or not text:
            continue
        if speaker.startswith("#") or len(speaker.split()) > 8:
            continue
        metadata = {
            "speaker": speaker,
            "sender_name": speaker,
            "persona_plus": True,
        }
        date_text = str(match.groupdict().get("date") or "").strip()
        if date_text:
            metadata["date"] = date_text
            metadata["sent_at"] = date_text
        return {"text": normalize_whitespace(f"{speaker}: {text}"), "metadata": metadata}
    return None


def _parse_intake_comment(body: str) -> dict[str, str]:
    metadata: dict[str, str] = {}
    for part in re.split(r"[|\n;]+", body or ""):
        if ":" in part:
            raw_key, raw_value = part.split(":", 1)
        elif "=" in part:
            raw_key, raw_value = part.split("=", 1)
        else:
            continue
        key = raw_key.strip().lower().replace("-", "_")
        value = raw_value.strip().strip("\"'")
        if key in PERSONA_INTAKE_KEYS and value:
            metadata[key] = value
    return metadata


def _persona_value_to_text(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, str):
        return normalize_whitespace(value)
    if isinstance(value, dict):
        if "text" in value:
            return _persona_value_to_text(value.get("text"))
        if "content" in value:
            return _persona_value_to_text(value.get("content"))
        parts = [_persona_value_to_text(child) for child in value.values()]
        return normalize_whitespace(" ".join(part for part in parts if part))
    if isinstance(value, list):
        parts = [_persona_value_to_text(item) for item in value]
        return normalize_whitespace(" ".join(part for part in parts if part))
    return normalize_whitespace(str(value))


def _clean_metadata(metadata: dict[str, Any]) -> dict[str, Any]:
    return {str(key): value for key, value in metadata.items() if value is not None and str(value).strip() != ""}


def _first_segment_date(segments: list[ExtractedSegment]) -> str | None:
    for segment in segments:
        metadata = dict(segment.metadata or {})
        for key in ("date", "sent_at", "approx_date"):
            value = str(metadata.get(key) or "").strip()
            if value:
                return value
    return None


def _persona_payload_title(payload: Any) -> str | None:
    if isinstance(payload, dict):
        for key in ("title", "name", "chat_name", "source"):
            value = str(payload.get(key) or "").strip()
            if value:
                return value[:80]
    return None


def _flatten_json(value: Any, path: list[str], lines: list[str]) -> None:
    if isinstance(value, dict):
        for key, child in value.items():
            key_string = str(key)
            if key_string.lower() in MEANINGLESS_JSON_KEYS:
                continue
            _flatten_json(child, [*path, key_string], lines)
        return
    if isinstance(value, list):
        for index, child in enumerate(value[:50]):
            _flatten_json(child, [*path, str(index)], lines)
        return
    if value is None:
        return
    if isinstance(value, str):
        text = normalize_whitespace(value)
        if not text:
            return
    else:
        text = str(value)
    key_path = ".".join(path)
    lines.append(f"{key_path}: {text}" if key_path else text)
